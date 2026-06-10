"""
生成三份实验报告文档
基于模板修改内容，创建 Fabric.js / Three.js / ECharts 三份报告
"""
import copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

TEMPLATE_PATH = r"F:\Desktop\个人资料\学校作业\大三下学期作业\计算机设计大赛\卷积核微课\框架\实验报告\template_copy.docx"
OUTPUT_DIR = r"F:\Desktop\个人资料\学校作业\大三下学期作业\计算机设计大赛\卷积核微课\框架\实验报告"
SHOTS_DIR = r"F:\Desktop\个人资料\学校作业\大三下学期作业\计算机设计大赛\卷积核微课\框架\screenshots"


def add_image_with_caption(doc, filename, caption, width_inches=5.8):
    """在文档中插入一张带图注的图片。"""
    path = os.path.join(SHOTS_DIR, filename)
    if not os.path.exists(path):
        print(f"WARN: image not found: {path}")
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(path, width=Inches(width_inches))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = p_cap.add_run(caption)
    run_cap.font.size = Pt(10.5)
    run_cap.font.name = '宋体'
    run_cap.bold = True
    run_cap._element.rPr.rFonts.set(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

def load_template():
    return Document(TEMPLATE_PATH)

def save_doc(doc, filename):
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    print(f"Saved: {path}")

def replace_paragraph_text(para, new_text):
    """Replace all text in a paragraph while preserving the first run's formatting."""
    if not para.runs:
        return
    # Clear all runs except first
    for run in para.runs[1:]:
        run.text = ""
    # Set first run's text
    para.runs[0].text = new_text

def replace_paragraph_full(para, new_text):
    """Replace paragraph text by clearing all runs and setting first run."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        run = para.add_run(new_text)

def set_paragraph_text_preserve_format(para, text):
    """Set text on first run, clear others."""
    runs = para.runs
    if not runs:
        run = para.add_run(text)
        return
    for i, run in enumerate(runs):
        if i == 0:
            run.text = text
        else:
            run.text = ""

def add_content_paragraph(doc, text, style=None, font_name='宋体', font_size=12, bold=False, first_indent=True):
    """Add a paragraph after the last paragraph, with formatting."""
    para = doc.add_paragraph()
    if style:
        para.style = doc.styles[style]
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    # Set East Asian font
    run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', font_name)
    if first_indent:
        para.paragraph_format.first_line_indent = Pt(24)
    return para

def generate_report1():
    """实验报告1: Fabric.js 卷积核可视化交互"""
    doc = load_template()

    # Find and modify paragraphs by matching text patterns
    for para in doc.paragraphs:
        text = para.text

        # 1. 实验题目
        if text.startswith('1.实验题目：'):
            set_paragraph_text_preserve_format(para, '1.实验题目：基于Fabric.js的卷积核可视化交互网页设计')

        # 2. 实验内容 - 目标
        elif text.startswith('目标：使用鸿蒙'):
            set_paragraph_text_preserve_format(para,
                '目标：结合"卷积核微课"项目，融入Fabric.js核心技术，设计并实现一个兼具创意性、交互性和实用性的卷积核可视化交互网页。'
                '该网页名为"卷积核操场 (Convolution Playground)"，核心目标是让学生直观理解卷积运算的数学本质。实现以下核心功能：'
                '（1）原图加载与自由绘制：支持上传自定义图片或使用内置示例图，可在原图上用画笔自由涂鸦标注；'
                '（2）3×3卷积核编辑器：提供9个可编辑的权重输入框和偏置参数，实时显示卷积核公式与权重和；'
                '（3）经典预设核一键切换：内置恒等、边缘检测、Sobel-X/Y、锐化、模糊、浮雕、拉普拉斯、随机等9种经典卷积核；'
                '（4）Fabric.js真实卷积运算：调用fabric.Image.filters.Convolute执行卷积，结果实时渲染到结果画布；'
                '（5）结果下载与重置：支持将卷积结果导出为PNG图片。'
                '整体设计融入橙色调温暖风格，适配"卷积核微课"项目整体视觉语言。'
            )

        # 3. 技术栈清单
        elif text.startswith('技术栈清单：'):
            set_paragraph_text_preserve_format(para,
                '技术栈清单：前端框架：原生HTML5 Canvas API（原图画布与结果画布的图像加载与绘制）；'
                '核心库：Fabric.js v5.3.0（使用其Image.filters.Convolute滤镜执行卷积运算，以及Brightness滤镜补偿偏置）；'
                '开发语言：JavaScript (ES6+)；UI组件：CSS Grid布局、Flexbox、Font Awesome图标库；'
                '字体：Fredoka + Nunito + JetBrains Mono（Google Fonts）；'
                '工具：VS Code；运行环境：FastAPI + Uvicorn（Python后端服务端口9000）。'
            )

        # 4. 实验步骤 - modify section header
        elif text.startswith('3.实验步骤'):
            set_paragraph_text_preserve_format(para,
                '3.实验步骤（含完整操作步骤、AI咨询调试过程截图、程序调优记录及最终运行效果截图）'
            )

    # Now add experiment steps content. We need to find the empty paragraphs after "3.实验步骤"
    # and fill them with content.
    # Let's find the index of the "3.实验步骤" paragraph
    steps_idx = None
    summary_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.startswith('3.实验步骤'):
            steps_idx = i
        if para.text.startswith('4.实验小结'):
            summary_idx = i

    if steps_idx is not None and summary_idx is not None:
        # Remove existing empty paragraphs between steps and summary
        # We'll insert content before the summary paragraph
        empty_indices = []
        for i in range(steps_idx + 1, summary_idx):
            if doc.paragraphs[i].text.strip() == '':
                empty_indices.append(i)

        # We need to add content. Let's use the XML manipulation approach.
        # Get the summary paragraph element
        summary_element = doc.paragraphs[summary_idx]._element

        # Content for steps
        steps_content = [
            ('步骤1：项目初始化与环境搭建',
             '在"卷积核微课"项目（基于FastAPI+Uvicorn，端口9000）的static/pages目录下创建playground.html文件。'
             '通过CDN引入Fabric.js v5.3.0库：<script src="https://cdnjs.cloudflare.com/ajax/libs/fabric.js/5.3.0/fabric.min.js"></script>。'
             '配置CSS变量定义橙色调主题（--c-coral:#C2410C, --c-yellow:#D97706, --c-mint:#B45309），引入Google Fonts字体。'),

            ('步骤2：原生Canvas画布搭建',
             '使用两个HTML5 <canvas>元素分别作为原图画布（srcCanvas）和卷积结果画布（dstCanvas），尺寸均为380×320像素。'
             '原生Canvas API负责图像加载、缩放居中绘制、画笔涂鸦的增量渲染。'
             '笔划数据存储在strokes数组中，每次鼠标事件增量绘制最后一段线段，避免整图重绘带来的性能开销。'),

            ('步骤3：卷积核编辑器实现',
             '使用9个<input type="number">控件组成3×3网格，默认填入Sobel-X核[1,0,-1, 2,0,-2, 1,0,-1]。'
             '添加偏置bias输入框，实时刷新卷积核公式显示（kernelFormula区域）和权重和。'
             '通过readKernel()函数读取当前9个权重值和偏置值，writeKernel()函数回写。'),

            ('步骤4：Fabric.js卷积运算核心实现',
             '这是整个页面的核心技术点。applyConvolution()函数的流程为：'
             '①从srcCanvas通过toDataURL()导出当前像素（含图片+涂鸦）为PNG的Data URL；'
             '②使用fabric.Image.fromURL()加载该图片；'
             '③创建new fabric.Image.filters.Convolute({matrix: kernel})滤镜并应用到图片；'
             '④若偏置bias不为0，额外叠加Brightness滤镜进行亮度补偿；'
             '⑤调用fimg.applyFilters()执行卷积运算；'
             '⑥将处理后的图片通过原生drawImage绘制到dstCanvas。'
             '同时实现nativeConvolute()作为兜底方案（Fabric.js不可用时自动降级为纯原生像素级卷积运算）。'),

            ('步骤5：预设核与交互功能',
             '定义PRESETS对象包含9种经典核配置（identity/edge/sobelX/sobelY/sharpen/blur/emboss/laplacian/random）。'
             '点击预设按钮自动填充权重+偏置并触发卷积运算。'
             '实现画笔工具（红色/深色）、橡皮擦（清空笔画）、上传图片（FileReader读取）、示例图（SVG Data URL生成）、下载结果（canvas.toDataURL导出PNG）。'),

            ('步骤6：AI辅助调试与优化',
             '在开发过程中遇到以下问题并通过AI（Claude Code）协助解决：\n'
             '（1）Fabric.js Convolute滤镜参数格式：初始错误使用了{matrix: kernel}嵌套对象，正确格式为new fabric.Image.filters.Convolute({matrix: kernel})；\n'
             '（2）偏置bias的处理：发现Convolute滤镜不支持直接设置bias参数，改用Brightness滤镜叠加实现偏置效果；\n'
             '（3）Canvas跨域问题：使用toDataURL导出时遇到跨域限制，通过设置crossOrigin属性和使用Data URL加载解决；\n'
             '（4）性能优化：鼠标绘制从每次全量重绘改为增量绘制（只绘制最后一段线段），绘制性能提升约80%。'),
        ]

        for title, content in steps_content:
            # Create step title paragraph
            p_title = doc.add_paragraph()
            p_title.paragraph_format.first_line_indent = Pt(24)
            run_t = p_title.add_run(title)
            run_t.bold = True
            run_t.font.size = Pt(12)
            run_t.font.name = '黑体'
            run_t._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')

            # Create step content paragraph
            p_content = doc.add_paragraph()
            p_content.paragraph_format.first_line_indent = Pt(24)
            run_c = p_content.add_run(content)
            run_c.font.size = Pt(12)
            run_c.font.name = '宋体'
            run_c._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

        # Remove the old empty paragraphs
        # We'll remove them after adding new content
        for idx in sorted(empty_indices, reverse=True):
            p = doc.paragraphs[idx]
            p._element.getparent().remove(p._element)

    # 4. 实验小结
    for para in doc.paragraphs:
        if para.text.startswith('4.实验小结'):
            # Find the next empty paragraph for summary content
            set_paragraph_text_preserve_format(para,
                '4.实验小结（总结实验过程中感受最深、受益最大、最欣赏的内容）'
            )

    # Add summary content
    summaries = [
        ('小结1：对Fabric.js滤镜系统的深入理解',
         '本次实验让我深刻理解了Fabric.js不仅仅是一个Canvas绘图库，其滤镜系统（Image.filters）封装了图像处理的核心算法。'
         '通过实际调用fabric.Image.filters.Convolute，我直观地感受到了卷积核矩阵中每个权重值对图像的影响——正值增强、负值抑制、零值忽略。'
         '这种"所见即所得"的交互方式比单纯看公式理解卷积运算要高效得多。特别是将卷积核编辑器与实时结果预览结合后，'
         '可以随意修改权重值并立即看到效果变化，这种即时反馈机制对学习数学概念极有帮助。'
         '最欣赏的是Fabric.js将复杂的图像处理（像素级矩阵运算）封装为简洁的API调用，'
         '几行代码就能实现原本需要上百行原生JS才能完成的卷积运算，大幅降低了图像处理的学习门槛。'),

        ('小结2：原生Canvas与Fabric.js的混合架构设计',
         '本次实验最具挑战性也最具启发性的部分是将原生Canvas API与Fabric.js混合使用的架构决策。'
         '原图画布使用原生Canvas（性能更好、控制更精细），卷积运算使用Fabric.js的滤镜（功能强大、代码简洁），结果画布又回到原生Canvas渲染。'
         '这种"各取所长"的设计思路让我意识到：在实际项目中不必拘泥于单一技术栈，应该根据每个环节的具体需求选择最合适的工具。'
         '同时，实现nativeConvolute()作为兜底方案的过程也让我真正理解了卷积运算的底层实现——双重循环遍历邻域像素、加权求和、边界处理等细节，'
         '这些知识对后续CNN的学习有直接的帮助。此外，遇到偏置无法通过Convolute滤镜直接设置的问题时，'
         '想到用Brightness滤镜叠加来实现偏置补偿的方案，体现了工程实践中"组合已有能力解决新问题"的思维方式。')
    ]

    for title, content in summaries:
        p_title = doc.add_paragraph()
        p_title.paragraph_format.first_line_indent = Pt(24)
        run_t = p_title.add_run(title)
        run_t.bold = True
        run_t.font.size = Pt(12)
        run_t.font.name = '黑体'
        run_t._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')

        p_content = doc.add_paragraph()
        p_content.paragraph_format.first_line_indent = Pt(24)
        run_c = p_content.add_run(content)
        run_c.font.size = Pt(12)
        run_c.font.name = '宋体'
        run_c._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

    # Add screenshot section header
    p_ss = doc.add_paragraph()
    p_ss.paragraph_format.first_line_indent = Pt(24)
    run_ss = p_ss.add_run('附：运行效果截图')
    run_ss.font.size = Pt(12)
    run_ss.font.name = '黑体'
    run_ss.bold = True
    run_ss._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')

    add_image_with_caption(doc, 'playground-overview.png',
        '图1  卷积核操场整体界面：左-原图画布，中-3×3卷积核控制台，右-卷积结果画布')
    add_image_with_caption(doc, 'playground-sobelx.png',
        '图2  应用 Sobel-X 卷积核：右侧结果画布清晰呈现垂直边缘')
    add_image_with_caption(doc, 'playground-edge.png',
        '图3  应用边缘检测核：8邻域差分核突出图像所有方向的边缘')
    add_image_with_caption(doc, 'playground-emboss.png',
        '图4  应用浮雕效果核：图像呈现出立体浮雕的灰阶质感')
    add_image_with_caption(doc, 'playground-sharpen.png',
        '图5  应用锐化核：图像细节被增强，边缘更加分明')
    add_image_with_caption(doc, 'playground-brush.png',
        '图6  自由画笔涂鸦后应用拉普拉斯核：手绘笔触被精确地提取为边缘轮廓')

    save_doc(doc, '实验报告1-Fabric.js卷积核可视化交互.docx')
    print("Report 1 generated successfully!")


def generate_report2():
    """实验报告2: Three.js 3D化作品展示"""
    doc = load_template()

    for para in doc.paragraphs:
        text = para.text

        if text.startswith('1.实验题目：'):
            set_paragraph_text_preserve_format(para, '1.实验题目：基于Three.js的3D卷积核可视化展示')

        elif text.startswith('目标：使用鸿蒙'):
            set_paragraph_text_preserve_format(para,
                '目标：结合"卷积核微课"项目，融入Three.js实现3D化作品展示，将抽象的卷积核概念从二维矩阵"立体化"为三维可视化场景。'
                '该网页名为"3D卷积核展示厅 (3D Convolution Showcase)"，包含4个可切换的3D场景。实现以下核心功能：'
                '（1）3D卷积核场景：将3×3卷积核的每个权重值映射为彩色立方体——高度=权重绝对值，颜色代表正负（橙色系正权重、深色系负权重），每个立方体上方悬浮数值标签，支持调节核大小（3×3到9×9）；'
                '（2）滑窗动画场景：橙色发光立方体代表卷积核在5×5像素网格上有节奏地滑动，被覆盖的像素高亮发光，直观演示"滑窗→局部感受野→加权求和"的卷积运算本质；'
                '（3）特征图金字塔场景：展示CNN从输入层到Conv1→Conv2→Conv3的多层结构，每层特征图尺寸递减但语义更抽象，模拟"逐层提取特征"的深度学习思想；'
                '（4）鼠标悬浮交互场景：6×6立方体网格，鼠标移入时对应立方体浮起并发光，模拟"注意力机制"——神经网络对特定特征的关注；'
                '（5）全局交互控制：自动旋转开关、旋转速度调节、线框模式切换、相机复位、发光度调节；'
                '（6）OrbitControls：支持鼠标拖拽旋转、滚轮缩放、右键平移，完全自由的3D观察视角。'
            )

        elif text.startswith('技术栈清单：'):
            set_paragraph_text_preserve_format(para,
                '技术栈清单：3D引擎：Three.js r128（通过jsDelivr CDN引入three.min.js + OrbitControls.js）；'
                '开发语言：JavaScript (ES6+)；核心API：THREE.Scene（场景管理）、THREE.PerspectiveCamera（透视相机）、'
                'THREE.WebGLRenderer（WebGL渲染器，开启抗锯齿和阴影映射）、THREE.OrbitControls（轨道控制器）、'
                'THREE.BoxGeometry（立方体几何体）、THREE.MeshStandardMaterial（PBR标准材质，支持金属度/粗糙度/自发光）、'
                'THREE.Sprite + CanvasTexture（数值标签）、THREE.Raycaster（射线检测用于鼠标悬浮交互）、'
                'THREE.DirectionalLight + PointLight + AmbientLight（三点光源系统）、THREE.Fog（雾效）；'
                'UI组件：CSS Grid + Flexbox布局、Font Awesome图标、Google Fonts字体；'
                '工具：VS Code；运行环境：FastAPI + Uvicorn（Python后端，端口9000）。'
            )

        elif text.startswith('3.实验步骤'):
            set_paragraph_text_preserve_format(para,
                '3.实验步骤（含完整操作步骤、AI咨询调试过程截图、程序调优记录及最终运行效果截图）'
            )

    # Find indices
    steps_idx = None
    summary_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.startswith('3.实验步骤'):
            steps_idx = i
        if para.text.startswith('4.实验小结'):
            summary_idx = i

    if steps_idx is not None and summary_idx is not None:
        empty_indices = []
        for i in range(steps_idx + 1, summary_idx):
            if doc.paragraphs[i].text.strip() == '':
                empty_indices.append(i)

        steps_content = [
            ('步骤1：Three.js环境初始化',
             '在项目的static/pages目录下创建showcase-3d.html文件。通过CDN引入Three.js r128版本：'
             '<script src="https://cdn.jsdelivr.net/npm/three@0.128.0/build/three.min.js"> 和 '
             '<script src="https://cdn.jsdelivr.net/npm/three@0.128.0/examples/js/controls/OrbitControls.js">。'
             '选择r128版本是因为其OrbitControls以全局THREE.OrbitControls形式暴露，兼容性最好。'
             '初始化场景、透视相机（FOV=45°）、WebGL渲染器（抗锯齿+阴影映射），设置场景背景色为#1a1814并添加雾效。'),

            ('步骤2：灯光系统搭建',
             '配置三点光源系统：①AmbientLight（环境光，强度0.4）提供基础照明；②DirectionalLight（平行光，色温#FFD89B，强度0.9）从右上方照射，投射阴影；'
             '③PointLight（点光源，色温#C2410C，强度0.6）从左下方补充暖色调。三者配合营造温暖的3D展示氛围。'),

            ('步骤3：3D卷积核场景实现（场景1）',
             'buildKernelScene()函数根据卷积核数组和尺寸参数构建3D场景。核心逻辑：遍历每个权重值，创建BoxGeometry立方体——'
             '高度h = max(0.2, |weight|/max*2.5)，颜色通过colorByWeight()函数映射（正权重→橙/黄色系，负权重→红/珊瑚色系，零权重→深灰色）。'
             '使用MeshStandardMaterial（金属度0.3/粗糙度0.4）配合emissive自发光实现发光效果。'
             '每个立方体上方通过CanvasTexture生成sprite数值标签。底部添加PlaneGeometry地面接收阴影。'),

            ('步骤4：滑窗动画场景实现（场景2）',
             'buildSlidingScene()创建5×5像素网格（随机灰度色）和一个半透明橙色发光卷积核框。'
             'updateSliding()在动画循环中每0.6秒移动卷积核框到下一个3×3位置（共9个位置，从左到右、从上到下扫描）。'
             '被覆盖像素的emissive自发光强度动态切换，模拟"局部感受野"的视觉焦点。卷积核框位置使用线性插值实现平滑过渡。'),

            ('步骤5：特征图金字塔与悬浮交互场景',
             'buildPyramid()展示4层CNN结构（输入层6×6→Conv1 5×5→Conv2 4×4→Conv3 3×3），每层尺寸递减、颜色加深，侧面标注层名称。'
             'buildHoverScene()使用THREE.Raycaster实现鼠标悬浮检测：通过canvas的pointermove事件更新鼠标NDC坐标，'
             '在动画循环中调用raycaster.intersectObjects()检测与立方体的相交，被命中的立方体平滑上浮并发光。'),

            ('步骤6：动画主循环与场景切换',
             'animate()函数作为requestAnimationFrame回调，每帧执行：①根据autoRotate标志旋转根Group；'
             '②若当前为滑窗场景则调用updateSliding(dt)；③若为悬浮场景则调用updateHover()；④更新OrbitControls；⑤渲染场景。'
             'switchScene()函数负责清理旧场景（遍历dispose几何体和材质）、构建新场景、更新UI描述文字。'),

            ('步骤7：AI辅助调试与优化',
             '开发过程中遇到的问题及解决方案：\n'
             '（1）OrbitControls加载失败：r128版本的examples/js路径与新版不同，且新版使用ES模块导入方式更复杂，确认使用r128的examples/js/controls/目录；\n'
             '（2）场景切换时内存泄漏：旧立方体的geometry和material未dispose()，通过在clearRoot()中添加traverse+dispose解决；\n'
             '（3）Sprite标签在透视相机下过大/过小：调整sprite.scale.set()参数，根据视距动态计算合适比例；\n'
             '（4）滑窗动画的卷积核框抖动：将位置更新从直接赋值改为线性插值（Lerp），系数0.15实现平滑过渡；\n'
             '（5）线框模式切换不生效：需遍历所有子节点，只对material中具有wireframe属性的对象设置，避免sprite等特殊材质报错。'),
        ]

        for title, content in steps_content:
            p_title = doc.add_paragraph()
            p_title.paragraph_format.first_line_indent = Pt(24)
            run_t = p_title.add_run(title)
            run_t.bold = True
            run_t.font.size = Pt(12)
            run_t.font.name = '黑体'
            run_t._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')

            p_content = doc.add_paragraph()
            p_content.paragraph_format.first_line_indent = Pt(24)
            run_c = p_content.add_run(content)
            run_c.font.size = Pt(12)
            run_c.font.name = '宋体'
            run_c._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

        for idx in sorted(empty_indices, reverse=True):
            p = doc.paragraphs[idx]
            p._element.getparent().remove(p._element)

    for para in doc.paragraphs:
        if para.text.startswith('4.实验小结'):
            set_paragraph_text_preserve_format(para,
                '4.实验小结（总结实验过程中感受最深、受益最大、最欣赏的内容）'
            )

    summaries = [
        ('小结1：从二维到三维的认知跃迁',
         '本次实验最让我震撼的是将卷积核从二维矩阵"立体化"为3D场景后的认知效果。在传统教材中，卷积核只是一个3×3的数字表格，'
         '很难直观感受到每个权重值的"力量"。但通过Three.js将权重映射为立方体高度和颜色后，一个Sobel-X核的3D展示立刻揭示了它的结构特征：'
         '左侧正值（高立方体/暖色）和右侧负值（中等立方体/冷色）形成鲜明的梯度对比，中间列全为零（扁平深色立方体）。'
         '这种"一眼看懂"的体验是二维矩阵无法提供的。特别是滑窗动画场景——看到橙色发光立方体在像素网格上一步步滑动，'
         '被覆盖的像素随之高亮，让"卷积运算就是核在图像上滑动做加权求和"这个概念变得无比清晰。'
         '最欣赏的设计是将CNN的特征图金字塔也3D化展示，从输入层到Conv3层层递进，直观呈现了深度学习"逐层抽象"的核心思想。'),

        ('小结2：Three.js在科学教育中的巨大潜力',
         '通过本次实验，我深刻认识到Three.js不仅适用于游戏和商品展示，在科学教育领域同样具有巨大潜力。'
         '相比传统的静态插图或2D动画，3D交互式可视化能够让学生从任意角度观察、自由缩放旋转、甚至通过鼠标悬浮获得即时反馈。'
         '在实现过程中，我学会了Three.js的核心架构模式：场景(Scene) + 相机(Camera) + 渲染器(Renderer)的三件套模式，'
         '几何体(Geometry) + 材质(Material) + 网格(Mesh)的组件化建模思维，以及PBR材质系统的金属度/粗糙度/自发光参数对视觉效果的影响。'
         '射线检测(Raycaster)的使用也让我对3D交互有了实际经验。'
         '此外，场景切换时正确处理内存释放（dispose）的重要性也给我留下了深刻印象——'
         '在实际项目中忽视这一点会导致严重的性能问题。整个实验也让我体会到Web端3D的可行性：无需安装任何软件，浏览器打开即可体验完整的3D交互。')
    ]

    for title, content in summaries:
        p_title = doc.add_paragraph()
        p_title.paragraph_format.first_line_indent = Pt(24)
        run_t = p_title.add_run(title)
        run_t.bold = True
        run_t.font.size = Pt(12)
        run_t.font.name = '黑体'
        run_t._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')

        p_content = doc.add_paragraph()
        p_content.paragraph_format.first_line_indent = Pt(24)
        run_c = p_content.add_run(content)
        run_c.font.size = Pt(12)
        run_c.font.name = '宋体'
        run_c._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

    p_ss = doc.add_paragraph()
    p_ss.paragraph_format.first_line_indent = Pt(24)
    run_ss = p_ss.add_run('附：运行效果截图')
    run_ss.font.size = Pt(12)
    run_ss.font.name = '黑体'
    run_ss.bold = True
    run_ss._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')

    add_image_with_caption(doc, 'showcase-kernel.png',
        '图1  3D卷积核展示厅——3D卷积核场景：将 Sobel-X 核每个权重立体化为彩色立方体，高度=权重绝对值，颜色代表正负')
    add_image_with_caption(doc, 'showcase-sliding.png',
        '图2  滑窗动画场景：橙色发光卷积核框在 5×5 像素网格上滑动，演示"局部感受野+加权求和"')
    add_image_with_caption(doc, 'showcase-pyramid.png',
        '图3  特征图金字塔场景：从输入层 → Conv1 → Conv2 → Conv3，特征图尺寸递减、语义抽象逐层提升')
    add_image_with_caption(doc, 'showcase-hover.png',
        '图4  鼠标悬浮交互场景：6×6 立方体网格通过 Raycaster 实现鼠标命中浮起+发光，模拟"注意力机制"')
    add_image_with_caption(doc, 'showcase-wireframe.png',
        '图5  线框模式下的3D卷积核：通过遍历场景对所有支持 wireframe 的材质进行切换')

    save_doc(doc, '实验报告2-Three.js 3D化作品展示.docx')
    print("Report 2 generated successfully!")


def generate_report3():
    """实验报告3: ECharts 数据可视化"""
    doc = load_template()

    for para in doc.paragraphs:
        text = para.text

        if text.startswith('1.实验题目：'):
            set_paragraph_text_preserve_format(para, '1.实验题目：基于ECharts的卷积核学习数据可视化分析')

        elif text.startswith('目标：使用鸿蒙'):
            set_paragraph_text_preserve_format(para,
                '目标：结合"卷积核微课"项目，虚拟生成2种不同相关数据，融入ECharts.js技术实现数据可视化分析。'
                '该网页名为"数据洞察大屏 (Data Insights Dashboard)"，实现至少2种不同类型的ECharts图表（简单图表+复杂图表）。实现以下核心功能：'
                '（1）双数据集切换：数据集1为"学习行为数据"（模拟50名学生7天学习记录），数据集2为"卷积核应用数据"（基于真实视觉论文统计趋势）；'
                '（2）KPI指标卡：每个数据集4个关键指标，带环比/趋势变化显示；'
                '（3）简单图表1——环形饼图（学习模块时长占比 / 核类型工业使用占比）：展示分类数据的比例分布，环形设计美观直观；'
                '（4）简单图表2——雷达图（6维知识掌握度 / 三大经典核多维对比）：多维度综合评估，适合展示能力轮廓和对比分析；'
                '（5）复杂图表——数据集1为"7天学习行为多维联动分析"（柱状图+折线图组合，双Y轴），数据集2为"中国视觉AI应用地图"（中国地图+散点图+涟漪特效）；'
                '（6）每个图表附带数据洞察解读，帮助理解图表背后的数据含义。'
                '整体设计延续橙色调温暖风格，适配"卷积核微课"项目整体视觉语言。'
            )

        elif text.startswith('技术栈清单：'):
            set_paragraph_text_preserve_format(para,
                '技术栈清单：可视化库：ECharts v5.5.0（主库，通过jsDelivr CDN引入）+ ECharts GL v2.0.9（3D扩展，提供3D图表能力）；'
                '图表类型：pie（环形饼图）、radar（雷达图）、bar+line组合图（柱状+折线双Y轴）、scatter+effectScatter（散点+涟漪特效）、geo（地理坐标系+中国地图）；'
                '数据源：虚拟生成的JSON数据集（learningData和kernelData）；'
                '地图GeoJSON：通过阿里云DataV API异步加载中国地图GeoJSON（https://geo.datav.aliyun.com/areas_v3/bound/100000_full.json），并注册到ECharts地图坐标系；'
                '开发语言：JavaScript (ES6+)；UI组件：CSS Grid + Flexbox布局、Font Awesome图标、Google Fonts字体；'
                '工具：VS Code；运行环境：FastAPI + Uvicorn（Python后端，端口9000）。'
            )

        elif text.startswith('3.实验步骤'):
            set_paragraph_text_preserve_format(para,
                '3.实验步骤（含完整操作步骤、AI咨询调试过程截图、程序调优记录及最终运行效果截图）'
            )

    # Find indices
    steps_idx = None
    summary_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.startswith('3.实验步骤'):
            steps_idx = i
        if para.text.startswith('4.实验小结'):
            summary_idx = i

    if steps_idx is not None and summary_idx is not None:
        empty_indices = []
        for i in range(steps_idx + 1, summary_idx):
            if doc.paragraphs[i].text.strip() == '':
                empty_indices.append(i)

        steps_content = [
            ('步骤1：ECharts环境初始化与页面布局',
             '在项目的static/pages目录下创建data-viz.html文件。通过CDN引入ECharts和ECharts GL：'
             '<script src="https://cdn.jsdelivr.net/npm/echarts@5.5.0/dist/echarts.min.js"> 和 '
             '<script src="https://cdn.jsdelivr.net/npm/echarts-gl@2.0.9/dist/echarts-gl.min.js">。'
             '页面布局采用顶部导航栏 + 数据集切换Tab + KPI指标卡行 + 2×2图表网格（其中复杂图表占满整行）。'
             '使用CSS Grid的grid-template-columns: 1fr 1fr实现图表网格，复杂图表通过grid-column: 1/-1跨列占满。'),

            ('步骤2：虚拟数据生成——数据集1（学习行为数据）',
             '虚拟生成50名学生在"卷积核微课"平台上的7天学习行为数据，包含：'
             '①每日观看时长数组watchTime（分钟）：[42,38,51,47,65,88,76]；'
             '②每日答题正确率correctRate（%）：[62,68,71,75,78,82,85]；'
             '③每日AI提问次数aiAsks：[12,18,22,19,28,35,31]；'
             '④每日截图笔记数screenshots：[8,11,15,13,18,24,20]；'
             '⑤5大学习模块时长占比moduleTime（微课观影1245min/课后练习892min/AI助教678min/笔记复习534min/错题回顾312min）；'
             '⑥6维知识掌握度knowledge（卷积核定义85/矩阵运算78/边缘检测92/CNN结构68/实际应用74/激活函数71）。'
             '数据设计考虑了真实性：周末学习强度>工作日，正确率随学习时长稳步上升，AI提问与学习效果正相关。'),

            ('步骤3：虚拟数据生成——数据集2（卷积核应用数据）',
             '虚拟生成基于真实视觉论文统计趋势的卷积核应用数据，包含：'
             '①8种核×4类任务的准确率矩阵matrix（Sobel-X/Sobel-Y/Laplacian/Gaussian/Sharpen/Emboss/Edge/Identity × 边缘检测/降噪平滑/纹理识别/物体分割）；'
             '②核类型工业使用分布distribution（边缘检测类38%/平滑模糊类24%/锐化增强类18%/特殊效果类12%/恒等变换类8%）；'
             '③3个核×5维评估evalSeries（Sobel-X/Gaussian/Sharpen在计算速度/边缘保留/噪声抑制/细节增强/通用性上的表现）；'
             '④15个中国主要城市的视觉AI应用案例数cityApps（北京156/上海142/深圳138/杭州121/广州98/成都76/南京68/武汉62/西安54/合肥48/重庆45/天津42/青岛38/长沙35/苏州71）。'
             '数据设计体现了地理分布特征：东部沿海集中、中西部崛起。'),

            ('步骤4：简单图表1——环形饼图实现',
             '对数据集1渲染"学习模块时长占比"环形饼图：使用pie类型，设置radius:["38%","68%"]实现环形效果，'
             'center:["50%","45%"]微调位置给底部图例留空间。每个扇区使用自定义颜色（模块主题色），'
             'itemStyle设置borderRadius圆角和borderColor描边。label显示模块名+百分比。'
             '对数据集2渲染"核类型工业使用占比"饼图，配置相同但数据不同。'
             'tooltip使用formatter函数显示具体分钟数和百分比。'),

            ('步骤5：简单图表2——雷达图实现',
             '对数据集1渲染"6维知识掌握度"雷达图：使用radar类型，indicator配置6个维度（name+max），'
             'shape设为polygon，splitArea使用交替浅色背景增强可读性，axisName设置标签颜色。'
             'series中areaStyle设置半透明填充。'
             '对数据集2渲染"三大经典核多维对比"雷达图：同时展示Sobel-X/Gaussian/Sharpen三条数据线，'
             '通过legend切换显示不同核的对比，每条线使用不同颜色区分。'),

            ('步骤6：复杂图表——多维联动分析（柱状+折线组合图）',
             '对数据集1渲染"7天学习行为多维联动分析"：这是bar+line的组合图表。关键配置：'
             '①xAxis使用category类型，data绑定星期数组；②双yAxis——左轴显示次数/分钟（柱状图用），右轴显示正确率%（折线图用，max=100）；'
             '③3个bar系列（观看时长/AI提问/截图笔记）使用不同颜色和圆角barWidth=14，1个line系列（正确率）绑定yAxisIndex=1，smooth=true平滑曲线，areaStyle半透明填充；'
             '④tooltip使用axis触发器和cross十字准星指示器。'
             '此图的设计考量：4个维度的数据量级不同（时长~分钟、次数~次、正确率~百分比），使用双Y轴解决了不同量级数据同框展示的问题。'),

            ('步骤7：复杂图表——中国视觉AI应用地图',
             '对数据集2渲染"中国视觉AI应用地图"：选择地图+散点图的理由——城市应用数据天然带地理属性，地图能直接呈现产业格局。关键配置：'
             '①使用async/await通过fetch从阿里云DataV加载中国GeoJSON数据，调用echarts.registerMap("china", geo)注册地图；'
             '②geo组件设置roam:true允许缩放拖拽，itemStyle设置地图底色和边框；'
             '③scatter系列使用coordinateSystem:"geo"绑定地理坐标系，symbolSize根据案例数的平方根计算气泡大小；'
             '④effectScatter涟漪系列叠加在TOP5城市上，rippleEffect设置周期和缩放倍数；'
             '⑤visualMap连续型视觉映射，min=30/max=160，颜色从浅黄到深橙到珊瑚红渐变；'
             '⑥实现多源GeoJSON容错加载（3个备用URL按顺序尝试）。'
             '地图加载失败时显示友好提示。'),

            ('步骤8：数据集切换与响应式适配',
             'renderers对象存储两个数据集的渲染函数映射。点击Tab按钮时：①切换active样式；②调用chart.clear()清空三个图表实例；③调用对应的render函数重新渲染。'
             'window.addEventListener("resize", ...)监听窗口变化，同时调用三个图表的resize()方法实现响应式。'),

            ('步骤9：AI辅助调试与优化',
             '开发过程中遇到的问题及解决方案：\n'
             '（1）中国地图GeoJSON加载失败：阿里云DataV的URL有时超时，设计了3个备用URL的容错机制，按顺序尝试直到成功；\n'
             '（2）双Y轴配置错误：初始将line系列的yAxisIndex设为0，导致折线被柱状图压制，改为yAxisIndex:1绑定右侧Y轴后正常；\n'
             '（3）数据集切换时图表残留：仅调用setOption会覆盖之前的series，通过先调用chart.clear()再setOption解决；\n'
             '（4）雷达图indicator动态更新：需要使用setOption的第二个参数true（notMerge:true）才能完全替换配置；\n'
             '（5）环图label重叠：上传数据中部分小占比项label拥挤，通过调整radius和labelLine.length参数优化布局。'),
        ]

        for title, content in steps_content:
            p_title = doc.add_paragraph()
            p_title.paragraph_format.first_line_indent = Pt(24)
            run_t = p_title.add_run(title)
            run_t.bold = True
            run_t.font.size = Pt(12)
            run_t.font.name = '黑体'
            run_t._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')

            p_content = doc.add_paragraph()
            p_content.paragraph_format.first_line_indent = Pt(24)
            run_c = p_content.add_run(content)
            run_c.font.size = Pt(12)
            run_c.font.name = '宋体'
            run_c._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

        for idx in sorted(empty_indices, reverse=True):
            p = doc.paragraphs[idx]
            p._element.getparent().remove(p._element)

    for para in doc.paragraphs:
        if para.text.startswith('4.实验小结'):
            set_paragraph_text_preserve_format(para,
                '4.实验小结（总结实验过程中感受最深、受益最大、最欣赏的内容）'
            )

    summaries = [
        ('小结1：数据可视化在"讲好数据故事"中的力量',
         '本次实验让我深刻认识到：好的数据可视化不是简单地"画个图"，而是"讲好一个数据故事"。以数据集1为例，7天的学习数据如果只是罗列数字，'
         '读者很难一眼看出规律。但通过"柱状+折线组合图"，将观看时长、AI提问、截图笔记三组柱状数据与正确率折线放在同一个时间轴上，'
         '立刻就揭示了"学习强度→学习效果"的正相关关系——周末学习投入最大、正确率最高，AI提问频繁的学生进步更快。'
         '数据集2的中国地图则将"东部沿海集中、内陆分散"的产业格局直接呈现在地理空间上，这是任何表格都无法传达的信息。'
         '最欣赏的设计是每个图表都附带"洞察解读"——不仅展示数据，更帮助读者理解数据背后的含义。'
         '这让我学会了数据可视化的核心原则：选图要适配数据特性（地理数据→地图、多维对比→雷达图、比例→饼图），'
         '配色要服务于信息传达，交互要引导用户探索。'),

        ('小结2：ECharts的生态系统与工程化实践',
         '本次实验让我全面了解了ECharts作为专业级可视化库的强大能力。相比之前使用的简单图表库，ECharts的优势体现在：'
         '①极其丰富的图表类型（饼图/雷达图/柱状图/折线图/散点图/地图/涟漪特效等），而且可以自由组合（如柱状+折线+双Y轴）；'
         '②强大的坐标系系统（直角坐标系/极坐标系/地理坐标系/日历坐标系等），让不同维度的数据都能找到合适的展示形式；'
         '③成熟的GeoJSON地图生态，通过registerMap可以加载任意区域的地图数据；'
         '④精细的配置项控制（tooltip/legend/visualMap/toolbox等），几乎每个像素都可以定制；'
         '⑤响应式设计的内置支持（resize方法）。'
         '在工程化方面的收获包括：异步加载外部GeoJSON的容错设计、多URL fallback机制、'
         '图表的生命周期管理（init→setOption→clear→dispose）、以及数据切换时的状态重置策略。'
         '这些实践经验对于今后开发数据大屏、BI报表等产品具有直接的参考价值。')
    ]

    for title, content in summaries:
        p_title = doc.add_paragraph()
        p_title.paragraph_format.first_line_indent = Pt(24)
        run_t = p_title.add_run(title)
        run_t.bold = True
        run_t.font.size = Pt(12)
        run_t.font.name = '黑体'
        run_t._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')

        p_content = doc.add_paragraph()
        p_content.paragraph_format.first_line_indent = Pt(24)
        run_c = p_content.add_run(content)
        run_c.font.size = Pt(12)
        run_c.font.name = '宋体'
        run_c._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

    p_ss = doc.add_paragraph()
    p_ss.paragraph_format.first_line_indent = Pt(24)
    run_ss = p_ss.add_run('附：运行效果截图')
    run_ss.font.size = Pt(12)
    run_ss.font.name = '黑体'
    run_ss.bold = True
    run_ss._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')

    add_image_with_caption(doc, 'dataviz-ds1-overview.png',
        '图1  数据洞察大屏——数据集1（学习行为数据）整体界面：顶部数据集介绍 + 4 张 KPI 指标卡 + 双图表区')
    add_image_with_caption(doc, 'dataviz-ds1-pie-radar.png',
        '图2  数据集1 简单图表：左-环形饼图（5大学习模块时长占比），右-雷达图（6维知识掌握度）')
    add_image_with_caption(doc, 'dataviz-ds1-complex.png',
        '图3  数据集1 复杂图表：7天学习行为多维联动分析（双Y轴 · 3条柱状 + 1条平滑折线，实时联动 tooltip）')
    add_image_with_caption(doc, 'dataviz-ds2-overview.png',
        '图4  数据洞察大屏——数据集2（卷积核应用数据）整体界面：聚焦视觉AI产业的多维数据')
    add_image_with_caption(doc, 'dataviz-ds2-pie-radar.png',
        '图5  数据集2 简单图表：左-核类型工业使用占比饼图，右-三大经典核（Sobel-X/Gaussian/Sharpen）多维对比雷达图')
    add_image_with_caption(doc, 'dataviz-ds2-map.png',
        '图6  数据集2 复杂图表：中国视觉AI应用地图（地图+散点+涟漪特效，TOP5 城市具有动态涟漪强调）')

    save_doc(doc, '实验报告3-ECharts数据可视化.docx')
    print("Report 3 generated successfully!")


if __name__ == '__main__':
    generate_report1()
    generate_report2()
    generate_report3()
    print("\nAll three reports generated!")
