"""
将截图嵌入到三份已生成的实验报告 docx 文件中。
策略：移除已有的纯文字"附：运行效果截图"段落（如果存在），
然后在文档末尾追加：截图小节标题 + 每张图片 + 图注。
"""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORTS_DIR = r"F:\Desktop\个人资料\学校作业\大三下学期作业\计算机设计大赛\卷积核微课\框架\实验报告"
SHOTS_DIR = r"F:\Desktop\个人资料\学校作业\大三下学期作业\计算机设计大赛\卷积核微课\框架\screenshots"


def add_image_with_caption(doc, filename, caption, width_inches=5.8):
    path = os.path.join(SHOTS_DIR, filename)
    if not os.path.exists(path):
        print(f"WARN: image not found: {path}")
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(path, width=Inches(width_inches))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = p_cap.add_run(caption)
    run_cap.font.size = Pt(10.5)
    run_cap.font.name = '宋体'
    run_cap.bold = True
    run_cap._element.rPr.rFonts.set(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')


def add_appendix_header(doc, text='附：运行效果截图'):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = '黑体'
    r.bold = True
    r._element.rPr.rFonts.set(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')


def remove_old_appendix(doc):
    """移除以"附：运行效果截图"开头的旧文字段落（含其后所有图1...图N行）。
    旧实现是单段落多行，匹配该段落直接删除即可。"""
    for para in list(doc.paragraphs):
        text = para.text.strip()
        if text.startswith('附：运行效果截图'):
            para._element.getparent().remove(para._element)


def process(filename, images):
    src = os.path.join(REPORTS_DIR, filename)
    if not os.path.exists(src):
        print(f"SKIP: {src} not found")
        return
    doc = Document(src)
    remove_old_appendix(doc)
    add_appendix_header(doc)
    for fname, caption in images:
        add_image_with_caption(doc, fname, caption)
    doc.save(src)
    print(f"Updated: {src}")


REPORT1 = ('实验报告1-Fabric.js卷积核可视化交互.docx', [
    ('playground-overview.png', '图1  卷积核操场整体界面：左-原图画布，中-3×3卷积核控制台，右-卷积结果画布'),
    ('playground-sobelx.png',   '图2  应用 Sobel-X 卷积核：右侧结果画布清晰呈现垂直边缘'),
    ('playground-edge.png',     '图3  应用边缘检测核：8邻域差分核突出图像所有方向的边缘'),
    ('playground-emboss.png',   '图4  应用浮雕效果核：图像呈现出立体浮雕的灰阶质感'),
    ('playground-sharpen.png',  '图5  应用锐化核：图像细节被增强，边缘更加分明'),
    ('playground-brush.png',    '图6  自由画笔涂鸦后应用拉普拉斯核：手绘笔触被精确地提取为边缘轮廓'),
])

REPORT2 = ('实验报告2-Three.js 3D化作品展示.docx', [
    ('showcase-kernel.png',    '图1  3D卷积核场景：将 Sobel-X 核每个权重立体化为彩色立方体（高度=权重绝对值，颜色代表正负）'),
    ('showcase-sliding.png',   '图2  滑窗动画场景：橙色发光卷积核框在 5×5 像素网格上滑动，演示"局部感受野+加权求和"'),
    ('showcase-pyramid.png',   '图3  特征图金字塔场景：从输入层 → Conv1 → Conv2 → Conv3，特征图尺寸递减、语义抽象逐层提升'),
    ('showcase-hover.png',     '图4  鼠标悬浮交互场景：6×6 立方体网格通过 Raycaster 实现鼠标命中浮起+发光，模拟"注意力机制"'),
    ('showcase-wireframe.png', '图5  线框模式下的3D卷积核：通过遍历场景对所有支持 wireframe 的材质进行切换'),
])

REPORT3 = ('实验报告3-ECharts数据可视化.docx', [
    ('dataviz-ds1-overview.png',  '图1  数据洞察大屏——数据集1（学习行为数据）整体界面：顶部数据集介绍 + 4 张 KPI 指标卡 + 双图表区'),
    ('dataviz-ds1-pie-radar.png', '图2  数据集1 简单图表：左-环形饼图（5大学习模块时长占比），右-雷达图（6维知识掌握度）'),
    ('dataviz-ds1-complex.png',   '图3  数据集1 复杂图表：7天学习行为多维联动分析（双Y轴 · 3条柱状 + 1条平滑折线）'),
    ('dataviz-ds2-overview.png',  '图4  数据洞察大屏——数据集2（卷积核应用数据）整体界面：聚焦视觉AI产业的多维数据'),
    ('dataviz-ds2-pie-radar.png', '图5  数据集2 简单图表：左-核类型工业使用占比饼图，右-三大经典核多维对比雷达图'),
    ('dataviz-ds2-map.png',       '图6  数据集2 复杂图表：中国视觉AI应用地图（地图+散点+涟漪特效，TOP5 城市动态涟漪强调）'),
])


if __name__ == '__main__':
    for filename, images in (REPORT1, REPORT2, REPORT3):
        process(filename, images)
    print("\nAll reports updated with images.")
